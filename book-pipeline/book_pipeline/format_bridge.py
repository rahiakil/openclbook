"""Import non-markdown sources into ``manuscript/draft.md`` and export merged staging to txt/docx/md."""

from __future__ import annotations

import json
import re
import shutil
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from book_pipeline.config import load_settings
from book_pipeline.gutenberg import strip_project_gutenberg_boilerplate
from book_pipeline.ingest import read_document


def _word_count(text: str) -> int:
    """Whitespace-delimited word count (same for draft and export bodies)."""
    if not text or not text.strip():
        return 0
    return len(text.split())


def _export_run_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _stamp_filename(name: str, stamp: str) -> str:
    """Insert ``_STAMP`` before the extension (``book.txt`` → ``book_20260418_120000.txt``)."""
    p = Path(name)
    if p.suffix:
        return f"{p.stem}_{stamp}{p.suffix}"
    return f"{name}_{stamp}"


@dataclass(frozen=True)
class StagingExportResult:
    """Paths and word counts from :func:`export_staging_merged`."""

    path: Path
    words_before: int | None
    words_after: int
    before_source: str | None
    staging_merged_words: int


def _md_to_plain_text(md: str) -> str:
    """Light markdown stripping for .txt export (no LLM)."""
    s = md
    s = re.sub(r"```[\s\S]*?```", "", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"^#{1,6}\s+", "", s, flags=re.MULTILINE)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip() + "\n"


def markdown_to_docx(md: str, out_path: Path) -> None:
    """Write a simple Word document from markdown-ish text (headings + paragraphs + local images)."""
    from docx import Document
    from docx.shared import Inches

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    ws = out_path.parent.parent.resolve()
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        s = line.strip()
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
        if m:
            alt = (m.group(1) or "").strip()
            pth = (m.group(2) or "").strip()
            # Allow workspace-relative paths like outputs/images/foo.png
            img_path = (ws / pth).resolve() if not Path(pth).is_absolute() else Path(pth)
            if img_path.is_file():
                if alt:
                    doc.add_paragraph(alt)
                try:
                    doc.add_picture(str(img_path), width=Inches(6.5))
                except Exception:
                    doc.add_paragraph(f"[image failed to embed: {pth}]")
            else:
                doc.add_paragraph(f"[missing image: {pth}]")
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph("• " + s[2:].strip())
        else:
            doc.add_paragraph(s)
    doc.save(str(out_path))


def import_source_to_draft(
    workspace: Path,
    input_path: Path,
    *,
    archive_sections: bool = False,
    requested_output_format: str = "md",
) -> Path:
    """
    Read ``.txt``, ``.md``, ``.docx`` (see ``read_document``), write ``manuscript/draft.md``.

    If ``archive_sections`` is True, existing ``sections/*.md`` are moved under
    ``.pipeline/archived_sections_<unix>/`` so the supervisor prefers the new draft.
    """
    ws = workspace.resolve()
    inp = input_path.expanduser().resolve()
    if not inp.is_file():
        raise FileNotFoundError(f"input not found: {inp}")

    text = strip_project_gutenberg_boilerplate(read_document(inp))
    settings = load_settings(ws, ws / "config.yaml")
    draft = ws / settings.manuscript_dir / "draft.md"
    draft.parent.mkdir(parents=True, exist_ok=True)

    if archive_sections:
        sec = ws / settings.sections_dir
        if sec.is_dir() and any(sec.glob("*.md")):
            dest = ws / ".pipeline" / f"archived_sections_{int(time.time())}"
            dest.mkdir(parents=True, exist_ok=True)
            for p in sorted(sec.glob("*.md")):
                shutil.move(str(p), str(dest / p.name))

    body = text.strip()
    if body and not body.lstrip().startswith("#"):
        title = inp.stem.replace("_", " ").strip() or "Manuscript"
        body = f"# {title}\n\n{body}\n"
    else:
        body = body + ("\n" if body and not body.endswith("\n") else "\n")

    draft.write_text(body, encoding="utf-8")

    meta = {
        "source_input": str(inp),
        "source_suffix": inp.suffix.lower(),
        "requested_output_format": requested_output_format.lower().strip(),
        "draft_written": str(draft.relative_to(ws)),
    }
    meta_dir = ws / ".pipeline"
    meta_dir.mkdir(parents=True, exist_ok=True)
    (meta_dir / "ingest_job.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return draft


def export_staging_merged(
    workspace: Path,
    output_format: str,
    *,
    output_name: str | None = None,
    stamp_filename: bool = True,
) -> StagingExportResult:
    """
    Read ``outputs/staging_merged.md`` and write:

    - ``md`` → ``outputs/staging_merged_export.md`` (or ``output_name``)
    - ``txt`` → ``outputs/staging_merged.txt`` (light markdown strip)
    - ``docx`` → ``outputs/staging_merged.docx``

    When ``stamp_filename`` is True (default), ``_YYYYMMDD_HHMMSS`` is inserted before the
    file extension so each run has a distinct artifact. Metadata is written to
    ``.pipeline/export_last.json`` including word counts for ``manuscript/draft.md`` (before)
    and the exported body (after). ``words_after`` for docx uses the merged markdown source
    (same material written into the document).
    """
    ws = workspace.resolve()
    settings = load_settings(ws, ws / "config.yaml")
    staging = ws / settings.outputs_dir / "staging_merged.md"
    if not staging.is_file():
        raise FileNotFoundError(f"missing merged output: {staging}")

    md = staging.read_text(encoding="utf-8", errors="replace")
    staging_merged_words = _word_count(md)
    fmt = output_format.lower().strip()
    out_dir = ws / settings.outputs_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    draft_path = ws / settings.manuscript_dir / "draft.md"
    words_before: int | None
    before_rel: str | None
    if draft_path.is_file():
        words_before = _word_count(draft_path.read_text(encoding="utf-8", errors="replace"))
        before_rel = str(draft_path.relative_to(ws))
    else:
        words_before = None
        before_rel = None

    stamp = _export_run_stamp() if stamp_filename else ""

    def _maybe_stamp(base_name: str) -> str:
        if not stamp_filename:
            return base_name
        return _stamp_filename(base_name, stamp)

    meta_dir = ws / ".pipeline"
    meta_dir.mkdir(parents=True, exist_ok=True)

    # Optional illustration appendix: produced by the plan-phase scene needs pass.
    scene_needs_md_path = out_dir / "scene_needs.md"
    appendix = ""
    if scene_needs_md_path.is_file():
        appendix = "\n\n---\n\n# Illustrations (auto)\n\n" + scene_needs_md_path.read_text(
            encoding="utf-8", errors="replace"
        ).strip() + "\n"

    if fmt == "md":
        name = output_name or "staging_merged_export.md"
        out = out_dir / _maybe_stamp(name)
        if not str(out.name).endswith(".md"):
            out = out.with_suffix(".md")
        body_after = md + appendix
        out.write_text(body_after, encoding="utf-8")
    elif fmt == "txt":
        name = output_name or "staging_merged.txt"
        out = out_dir / _maybe_stamp(name)
        if not str(out.name).endswith(".txt"):
            out = out.with_suffix(".txt")
        body_after = _md_to_plain_text(md + appendix)
        out.write_text(body_after, encoding="utf-8")
    elif fmt == "docx":
        name = output_name or "staging_merged.docx"
        out = out_dir / _maybe_stamp(name)
        if not str(out.name).endswith(".docx"):
            out = out.with_suffix(".docx")
        body_after = md + appendix
        markdown_to_docx(body_after, out)
    else:
        raise ValueError(f"unsupported output_format {output_format!r}; use md, txt, or docx")

    words_after = _word_count(body_after)
    export_meta = {
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "export_path": str(out.relative_to(ws)),
        "output_format": fmt,
        "stamp_filename": stamp_filename,
        "word_count_before": words_before,
        "word_count_before_source": before_rel,
        "word_count_after": words_after,
        "word_count_staging_merged_md": staging_merged_words,
    }
    (meta_dir / "export_last.json").write_text(json.dumps(export_meta, indent=2), encoding="utf-8")

    return StagingExportResult(
        path=out,
        words_before=words_before,
        words_after=words_after,
        before_source=before_rel,
        staging_merged_words=staging_merged_words,
    )
